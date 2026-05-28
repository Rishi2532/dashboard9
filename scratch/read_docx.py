import docx

doc = docx.Document("uploads/first page.docx")
print("=== PARAGRAPHS ===")
for p in doc.paragraphs:
    if p.text.strip():
        print(p.text)

print("=== TABLES ===")
for t in doc.tables:
    for r in t.rows:
        row_text = [cell.text.strip() for cell in r.cells]
        print(" | ".join(row_text))

print("=== SECTIONS (FOOTER/HEADER) ===")
for section in doc.sections:
    footer = section.footer
    for p in footer.paragraphs:
        if p.text.strip():
            print("Footer:", p.text)
